"""build_vancouver_replacement_fixture.py — PMC OA 1 論文から Vancouver 35-ref fixture を組成する.

Usage:
    uv run python tools/build_vancouver_replacement_fixture.py \\
        --pmc-id PMC13179246 \\
        --output tests/fixtures/vancouver_35refs/input_References.docx

実行手順:
    1. NCBI E-utilities efetch で 1 論文の PMC XML を取得 (.cache/pmc_xml/ にキャッシュ)
    2. XML の <ref-list> から全件 (最大 limit=999) を抽出
    3. JATS の <element-citation> / <mixed-citation> を Vancouver/AMA plain text に
       structured-field 経路で再組成
       Vancouver format: "Author A, Author B, et al. Title. Journal. Year;Vol(Iss):pp-pp.
       doi:10.xxxx/xxx"
    4. python-docx で番号付き段落 (1. ~ 35.) として docx 生成

依存: urllib (標準) / python-docx / xml.etree.ElementTree (標準)

Day23 Phase 5b. 旧 vancouver_24refs (OneDrive 由来、査読守秘義務懸念) の代替として
PMC OA CC BY 4.0 由来の fixture を構築. build_apa_fixture.py を template に改変.

Source: PMC13179246 — Supportive Care in Cancer (Springer Nature), 2026, CC BY 4.0
DOI: 10.1007/s00520-026-10782-z
"""
from __future__ import annotations

import argparse
import os
import re
import ssl
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
from pathlib import Path

PMC_EFETCH_URL = (
    "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
)
# Anchor cache to repo root so .gitignore's `.cache/` entry catches it
# regardless of CWD at invocation time.
REPO_ROOT = Path(__file__).resolve().parent.parent
CACHE_DIR = REPO_ROOT / ".cache" / "pmc_xml"
TIMEOUT_SECONDS = 30
# NCBI rate limit: with API key 10 req/s, without 3 req/s.
SLEEP_BETWEEN_FETCHES = 0.34


# --------------------------------------------------------------------------- #
# PMC XML fetch
# --------------------------------------------------------------------------- #
def fetch_pmc_xml(pmc_id: str, api_key: str | None = None) -> str:
    """Fetch PMC OA full-text XML via NCBI E-utilities efetch.

    Caches result to ``.cache/pmc_xml/{pmc_id}.xml`` to avoid re-fetching during
    iterative development.

    Args:
        pmc_id: PMC ID (e.g. ``"PMC13179246"``). The ``PMC`` prefix is stripped.
        api_key: optional NCBI API key.

    Returns:
        Raw XML response body as str (utf-8 decoded).
    """
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_path = CACHE_DIR / f"{pmc_id}.xml"
    if cache_path.exists():
        return cache_path.read_text(encoding="utf-8")

    numeric_id = pmc_id.removeprefix("PMC")
    params: dict[str, str] = {
        "db": "pmc",
        "id": numeric_id,
        "rettype": "xml",
    }
    if api_key:
        params["api_key"] = api_key
    url = PMC_EFETCH_URL + "?" + urllib.parse.urlencode(params)

    # Use certifi SSL context if available (avoids self-signed cert errors on macOS).
    try:
        import certifi
        ssl_ctx: ssl.SSLContext | None = ssl.create_default_context(
            cafile=certifi.where()
        )
    except ImportError:
        ssl_ctx = None

    with urllib.request.urlopen(url, timeout=TIMEOUT_SECONDS, context=ssl_ctx) as resp:
        body = resp.read().decode("utf-8")

    cache_path.write_text(body, encoding="utf-8")
    time.sleep(SLEEP_BETWEEN_FETCHES)
    return body


# --------------------------------------------------------------------------- #
# Reference extraction
# --------------------------------------------------------------------------- #
def extract_references(xml: str, limit: int = 999) -> list[dict[str, str]]:
    """Extract all (up to ``limit``) successfully-formatted references from PMC XML.

    Args:
        xml: PMC OA full-text XML body.
        limit: maximum number of references to extract (default: 999 = all).

    Returns:
        List of dicts with shape ``{"label": "1", "raw": "Smith A, et al. ..."}``.
    """
    root = ET.fromstring(xml)
    refs: list[dict[str, str]] = []
    for ref_el in root.iter("ref"):
        if len(refs) >= limit:
            break
        # Some publishers wrap citations in <citation-alternatives>;
        # use iter() for namespace-agnostic deep search, preferring element-citation.
        cit_el: ET.Element | None = None
        for candidate_tag in ("element-citation", "mixed-citation"):
            found = next(ref_el.iter(candidate_tag), None)
            if found is not None:
                cit_el = found
                break
        if cit_el is None:
            continue
        vancouver_text = format_as_vancouver(cit_el)
        if not vancouver_text:
            continue
        refs.append({"label": str(len(refs) + 1), "raw": vancouver_text})
    return refs


# --------------------------------------------------------------------------- #
# Vancouver formatting from JATS structured fields
# --------------------------------------------------------------------------- #
def format_as_vancouver(cit_el: ET.Element) -> str:
    """Convert a JATS citation element to Vancouver/ICMJE plain text.

    Vancouver format:
        Author A, Author B, Author C, et al. Title. Journal. Year;Vol(Iss):pp-pp.
        doi:10.xxxx/xxx

    - Surname + initials without periods (e.g. "Smith JA" not "Smith, J. A.")
    - Up to 6 authors, then "et al."
    - Journal name as provided in JATS (may be abbreviated)
    - Year;Vol(Iss):fpage-lpage
    - doi: prefix (not URL form)

    Returns empty string when both authors AND year are missing.
    """
    authors = _format_authors_vancouver(cit_el)
    year_el = cit_el.find("year")
    year = (year_el.text or "").strip() if year_el is not None else ""

    title_el = cit_el.find("article-title")
    if title_el is None:
        title_el = cit_el.find("chapter-title")
    title = _full_text(title_el).strip() if title_el is not None else ""

    source_el = cit_el.find("source")
    journal = _full_text(source_el).strip() if source_el is not None else ""

    def _t(tag: str) -> str:
        el = cit_el.find(tag)
        return (el.text or "").strip() if el is not None and el.text else ""

    volume = _t("volume")
    issue = _t("issue")
    fpage = _t("fpage")
    lpage = _t("lpage")

    doi = ""
    for pub_id in cit_el.findall("pub-id"):
        if pub_id.get("pub-id-type") == "doi" and pub_id.text:
            doi = pub_id.text.strip()
            break

    if not authors and not year:
        return ""

    parts: list[str] = []

    # Authors block
    if authors:
        author_part = authors
        if not author_part.endswith("."):
            author_part += "."
        parts.append(author_part)

    # Title
    if title:
        t = title.rstrip()
        if not t.endswith((".", "?", "!")):
            t = t + "."
        parts.append(t)

    # Journal;Vol(Iss):pages. Year
    if journal or year:
        jbits = journal if journal else ""
        if year:
            jbits = (jbits + ". " if jbits else "") + year
            if volume:
                jbits += f";{volume}"
                if issue:
                    jbits += f"({issue})"
        if fpage:
            pages = fpage if not lpage else f"{fpage}-{lpage}"
            jbits += f":{pages}"
        if not jbits.endswith("."):
            jbits += "."
        parts.append(jbits)

    body = " ".join(p for p in parts if p)

    if doi:
        body = body + f" doi:{doi}"

    return _collapse_whitespace(body)


# --------------------------------------------------------------------------- #
# Author handling (Vancouver style)
# --------------------------------------------------------------------------- #
def _format_authors_vancouver(cit_el: ET.Element) -> str:
    """Format authors per Vancouver/ICMJE from JATS name nodes.

    Vancouver format:
        - "Surname AB" (no periods, no comma between surname and initials)
        - Up to 6 authors listed; 7+ → first 6 + ", et al"
        - Separated by ", "
        - Organizational authors: as-is

    Args:
        cit_el: JATS citation element.

    Returns:
        Formatted author string, e.g. "Smith JA, Brown KL, et al"
    """
    name_nodes: list[ET.Element] = []
    for el in cit_el.iter():
        if el is cit_el:
            continue
        if el.tag in ("name", "string-name", "collab"):
            name_nodes.append(el)

    formatted: list[str] = []
    for n in name_nodes:
        if n.tag == "collab":
            collab_text = "".join(n.itertext()).strip()
            if collab_text:
                formatted.append(collab_text)
            continue
        surname_el = n.find("surname")
        given_el = n.find("given-names")
        surname = (surname_el.text or "").strip() if surname_el is not None else ""
        given = (given_el.text or "").strip() if given_el is not None else ""
        if not surname and not given:
            txt = (n.text or "").strip()
            if not txt:
                continue
            formatted.append(txt)
            continue
        initials = _initials_vancouver(given)
        if surname and initials:
            formatted.append(f"{surname} {initials}")
        elif surname:
            formatted.append(surname)
        elif initials:
            formatted.append(initials)

    if not formatted:
        return ""

    # Vancouver: up to 6 authors, then et al.
    MAX_AUTHORS = 6
    if len(formatted) > MAX_AUTHORS:
        return ", ".join(formatted[:MAX_AUTHORS]) + ", et al"
    return ", ".join(formatted)


def _initials_vancouver(given: str) -> str:
    """Convert given names to Vancouver initials (no periods).

    Examples:
        'John Anthony' -> 'JA'
        'J. A.'        -> 'JA'
        'JA'           -> 'JA'
        'K'            -> 'K'
    """
    given = given.strip()
    if not given:
        return ""
    # Remove dots
    given_nodot = given.replace(".", "")
    # Split by whitespace
    tokens = given_nodot.split()
    if len(tokens) > 1:
        # 'John Anthony' or 'J A' style → take first char of each
        return "".join(t[0].upper() for t in tokens if t)
    # Single token: could be 'JA', 'J', 'John'
    tok = tokens[0]
    if tok.isupper() and len(tok) >= 2:
        # Already compressed initials like 'JA', 'ABC'
        return tok
    # Single char or mixed case (like 'John') → take first char
    return tok[0].upper()


# --------------------------------------------------------------------------- #
# XML text helpers
# --------------------------------------------------------------------------- #
def _full_text(el: ET.Element | None) -> str:
    """Return concatenated text content (including child tags) of an element."""
    if el is None:
        return ""
    return "".join(el.itertext())


def _collapse_whitespace(s: str) -> str:
    """Normalize all whitespace runs (incl. newlines/tabs) to single spaces."""
    return re.sub(r"\s+", " ", s).strip()


# --------------------------------------------------------------------------- #
# DOCX assembly
# --------------------------------------------------------------------------- #
def build_docx(refs: list[dict[str, str]], output_path: Path) -> None:
    """Assemble numbered DOCX via python-docx.

    Structure:
        Heading 1: "References"
        Paragraph "1. {ref0.raw}"
        ...
    """
    from docx import Document  # local import: optional dependency

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("References", level=1)
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {ref['raw']}")
    doc.save(str(output_path))


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def _parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Build a Vancouver-style fixture DOCX from 1 PMC OA paper "
            "(Day23 Phase 5b). Source: PMC13179246 (Supportive Care in Cancer)."
        ),
    )
    parser.add_argument(
        "--pmc-id",
        default="PMC13179246",
        help="PMC ID for the source paper (default: PMC13179246).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("tests/fixtures/vancouver_35refs/input_References.docx"),
        help="Output DOCX path.",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=999,
        help="Maximum number of references to extract (default: 999 = all).",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = _parse_args(argv if argv is not None else sys.argv[1:])

    api_key = os.environ.get("NCBI_API_KEY") or None
    if not api_key:
        print(
            "WARN: NCBI_API_KEY not set; falling back to anonymous "
            "E-utilities rate limit (3 req/s).",
            file=sys.stderr,
        )

    pmc_id = args.pmc_id
    try:
        xml = fetch_pmc_xml(pmc_id, api_key=api_key)
    except (urllib.error.HTTPError, urllib.error.URLError,
            TimeoutError, OSError) as e:
        print(f"ERROR: fetch_pmc_xml({pmc_id}) failed: {e}", file=sys.stderr)
        return 2

    refs = extract_references(xml, limit=args.limit)
    print(f"[{pmc_id}] extracted {len(refs)} refs")

    build_docx(refs, args.output)
    print(f"wrote {len(refs)} refs to {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
