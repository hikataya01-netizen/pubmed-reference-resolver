"""build_apa_fixture.py — PMC OA 3 論文から APA 45-ref fixture を組成する.

Usage:
    uv run python tools/build_apa_fixture.py \\
        --psychology PMC11601046 \\
        --public-health PMC11404860 \\
        --psychology-religion PMC11165362 \\
        --output tests/fixtures/apa_45refs/input_References.docx \\
        --refs-per-paper 15

実行手順:
    1. NCBI E-utilities efetch で 3 論文の PMC XML を取得 (.cache/pmc_xml/ にキャッシュ)
    2. 各 XML の <ref-list> から先頭 N 件を抽出
    3. JATS の <element-citation> / <mixed-citation> を APA 7 plain text に
       structured-field 経路で常に再組成 (Frontiers の <name> 形式にも対応)
    4. python-docx で番号付き段落 (1. ~ 45.) として統合 docx 生成

依存: urllib (標準) / python-docx / xml.etree.ElementTree (標準)

Day16 Phase 0a (Task 1 implementation). 再現性確保のため commit する.
"""
from __future__ import annotations

import argparse
import os
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
from pathlib import Path

PMC_EFETCH_URL = (
    "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
)
# Anchor cache to repo root so .gitignore's `.cache/` entry catches it
# regardless of CWD at invocation time.
REPO_ROOT = Path(__file__).resolve().parent.parent
CACHE_DIR = REPO_ROOT / ".cache" / "pmc_xml"
TIMEOUT_SECONDS = 30
# NCBI rate limit: with API key 10 req/s, without 3 req/s.
# 0.34s sleep keeps us under 3 req/s even without a key.
SLEEP_BETWEEN_FETCHES = 0.34


# --------------------------------------------------------------------------- #
# PMC XML fetch
# --------------------------------------------------------------------------- #
def fetch_pmc_xml(pmc_id: str, api_key: str | None = None) -> str:
    """Fetch PMC OA full-text XML via NCBI E-utilities efetch.

    Caches result to ``.cache/pmc_xml/{pmc_id}.xml`` to avoid re-fetching during
    iterative development. After a live fetch, sleeps ``SLEEP_BETWEEN_FETCHES``
    seconds to stay under NCBI rate limits.

    Args:
        pmc_id: PMC ID (e.g. ``"PMC11601046"``). The ``PMC`` prefix is stripped
            for the ``id`` query param per NCBI convention.
        api_key: optional NCBI API key (raises rate limit to 10 req/s).

    Returns:
        Raw XML response body as str (utf-8 decoded).
    """
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_path = CACHE_DIR / f"{pmc_id}.xml"
    if cache_path.exists():
        return cache_path.read_text(encoding="utf-8")

    # NCBI expects the numeric ID without the "PMC" prefix.
    numeric_id = pmc_id.removeprefix("PMC")
    params: dict[str, str] = {
        "db": "pmc",
        "id": numeric_id,
        "rettype": "xml",
    }
    if api_key:
        params["api_key"] = api_key
    url = PMC_EFETCH_URL + "?" + urllib.parse.urlencode(params)

    with urllib.request.urlopen(url, timeout=TIMEOUT_SECONDS) as resp:
        body = resp.read().decode("utf-8")

    cache_path.write_text(body, encoding="utf-8")
    time.sleep(SLEEP_BETWEEN_FETCHES)
    return body


# --------------------------------------------------------------------------- #
# Reference extraction
# --------------------------------------------------------------------------- #
def extract_references(xml: str, limit: int = 15) -> list[dict[str, str]]:
    """Extract first ``limit`` successfully-formatted references from PMC XML.

    Uses ``root.iter("ref")`` for namespace-agnostic search (JATS variants
    differ in default-namespace usage across publishers).

    Counter logic: only successful (non-empty) APA 7 conversions count
    toward ``limit``. Failed conversions are silently skipped, ensuring the
    output always reaches ``limit`` entries when enough <ref> nodes exist.

    Args:
        xml: PMC OA full-text XML body.
        limit: number of references to extract per paper.

    Returns:
        List of dicts with shape ``{"label": "1", "raw": "Smith, J. ..."}``.
    """
    root = ET.fromstring(xml)
    refs: list[dict[str, str]] = []
    for ref_el in root.iter("ref"):
        if len(refs) >= limit:
            break
        cit_el = ref_el.find("element-citation")
        if cit_el is None:
            cit_el = ref_el.find("mixed-citation")
        if cit_el is None:
            continue
        apa_text = format_as_apa7(cit_el)
        if not apa_text:
            continue
        refs.append({"label": str(len(refs) + 1), "raw": apa_text})
    return refs


# --------------------------------------------------------------------------- #
# APA 7 formatting from JATS structured fields
# --------------------------------------------------------------------------- #
def format_as_apa7(cit_el: ET.Element) -> str:
    """Convert a JATS citation element to APA 7 plain text.

    CRITICAL: Always uses structured-field extraction regardless of element
    type (``element-citation`` vs ``mixed-citation``). This ensures clean APA 7
    output even from journals whose XML lacks tail-commas between ``<name>``
    children (e.g. Frontiers).

    Returns empty string when both authors AND year are missing — such a
    citation cannot be meaningfully rendered as APA 7.
    """
    authors = _format_authors(cit_el)
    year_el = cit_el.find("year")
    year = (year_el.text or "").strip() if year_el is not None else ""

    title_el = cit_el.find("article-title")
    if title_el is None:
        title_el = cit_el.find("chapter-title")
    title = _full_text(title_el).strip() if title_el is not None else ""

    source_el = cit_el.find("source")
    journal = _full_text(source_el).strip() if source_el is not None else ""

    def _t(tag: str) -> str:
        el = cit_el.find(tag)
        return (el.text or "").strip() if el is not None and el.text else ""

    volume = _t("volume")
    issue = _t("issue")
    fpage = _t("fpage")
    lpage = _t("lpage")

    doi = ""
    for pub_id in cit_el.findall("pub-id"):
        if pub_id.get("pub-id-type") == "doi" and pub_id.text:
            doi = pub_id.text.strip()
            break

    if not authors and not year:
        return ""

    # Build: "{authors} ({year}). {title}."
    parts: list[str] = []
    head = ""
    if authors:
        head = authors
    if year:
        head = (head + " " if head else "") + f"({year})."
    elif head:
        head = head + "."
    parts.append(head.strip())

    if title:
        # Ensure title ends with a period (preserve trailing ? or ! if present).
        t = title.rstrip()
        if not t.endswith((".", "?", "!")):
            t = t + "."
        parts.append(t)

    # Journal block: "{journal}, {volume}({issue}), {fpage}-{lpage}."
    # Preserve any trailing period in the journal name (e.g. "Front. Psychol.")
    # by inserting punctuation only when we have more to append.
    if journal:
        jbits = journal
        if volume:
            jbits += f", {volume}"
            if issue:
                jbits += f"({issue})"
        if fpage:
            pages = fpage if not lpage else f"{fpage}-{lpage}"
            jbits += f", {pages}"
        if not jbits.endswith("."):
            jbits += "."
        parts.append(jbits)

    body = " ".join(p for p in parts if p)

    if doi:
        body = body + f" https://doi.org/{doi}"

    return _collapse_whitespace(body)


# --------------------------------------------------------------------------- #
# Author handling
# --------------------------------------------------------------------------- #
def _format_authors(cit_el: ET.Element) -> str:
    """Format authors per APA 7 from ``<name>``, ``<string-name>``, ``<collab>`` nodes.

    Walks ``cit_el`` in document order, picking up ``<name>``, ``<string-name>``
    and ``<collab>`` (organizational author) children. Handles three given-name
    forms:
        - 'John Anthony'  -> 'J. A.'
        - 'JE'            -> 'J. E.'  (compressed initials, all-caps)
        - 'K. G.'         -> 'K. G.'  (pre-formatted, preserved)

    APA 7 author-list combiners:
        0          -> ''
        1          -> 'Smith, J. A.'
        2          -> 'Smith, J. A., & Brown, K. L.'
        3+         -> 'Smith, A., Brown, K., & Jones, M.'

    Organizational authors (``<collab>``) appear as-is per APA 7:
        'Department of Health and Social Care'
        'World Health Organization'
    """
    # Document-order walk: collect <name>, <string-name>, and <collab> nodes.
    name_nodes: list[ET.Element] = []
    for el in cit_el.iter():
        if el is cit_el:
            continue
        if el.tag in ("name", "string-name", "collab"):
            name_nodes.append(el)

    formatted: list[str] = []
    for n in name_nodes:
        if n.tag == "collab":
            # Organizational author: use full text content (strip child tags).
            collab_text = "".join(n.itertext()).strip()
            if collab_text:
                formatted.append(collab_text)
            continue
        surname_el = n.find("surname")
        given_el = n.find("given-names")
        surname = (surname_el.text or "").strip() if surname_el is not None else ""
        given = (given_el.text or "").strip() if given_el is not None else ""
        if not surname and not given:
            # Some <string-name> elements carry the whole name as text only.
            txt = (n.text or "").strip()
            if not txt:
                continue
            formatted.append(txt)
            continue
        initials = _normalize_initials(given)
        if surname and initials:
            formatted.append(f"{surname}, {initials}")
        elif surname:
            formatted.append(surname)
        elif initials:
            formatted.append(initials)

    if not formatted:
        return ""
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        return f"{formatted[0]}, & {formatted[1]}"
    return ", ".join(formatted[:-1]) + f", & {formatted[-1]}"


def _normalize_initials(given: str) -> str:
    """'John Anthony' -> 'J. A.'  /  'JE' -> 'J. E.'  /  'K. G.' -> 'K. G.'"""
    given = given.strip()
    if not given:
        return ""
    # If contains dots, preserve as-is (already formatted)
    if "." in given:
        return given.rstrip(".") + "."  # ensure trailing dot
    # Split by whitespace first
    tokens = given.split()
    if len(tokens) > 1:
        # 'John Anthony' style
        return ". ".join(t[0].upper() for t in tokens if t) + "."
    # Single token: could be 'J' or 'JE' (compressed)
    tok = tokens[0]
    if tok.isupper() and len(tok) >= 2:
        # 'JE' -> 'J. E.'  (treat each capital as separate initial)
        caps = re.findall(r"[A-Z]", tok)
        return ". ".join(caps) + "."
    # Single char or lowercase
    return tok[0].upper() + "."


# --------------------------------------------------------------------------- #
# XML text helpers
# --------------------------------------------------------------------------- #
def _full_text(el: ET.Element | None) -> str:
    """Return concatenated text content (including child tags) of an element."""
    if el is None:
        return ""
    return "".join(el.itertext())


def _collapse_whitespace(s: str) -> str:
    """Normalize all whitespace runs (incl. newlines/tabs) to single spaces."""
    return re.sub(r"\s+", " ", s).strip()


# --------------------------------------------------------------------------- #
# DOCX assembly
# --------------------------------------------------------------------------- #
def build_docx(refs: list[dict[str, str]], output_path: Path) -> None:
    """Assemble 45-ref numbered DOCX via python-docx.

    Structure:
        Heading 1: "References"
        Paragraph "1. {ref0.raw}"
        Paragraph "2. {ref1.raw}"
        ...
    """
    from docx import Document  # local import: optional dependency

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("References", level=1)
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {ref['raw']}")
    doc.save(str(output_path))


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def _parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Build a 45-reference APA fixture DOCX from 3 PMC OA papers "
            "(Day16 Phase 0a)."
        ),
    )
    parser.add_argument(
        "--psychology",
        required=True,
        help="PMC ID for the psychology paper (e.g. PMC11601046).",
    )
    parser.add_argument(
        "--public-health",
        required=True,
        help="PMC ID for the public-health paper (e.g. PMC11404860).",
    )
    parser.add_argument(
        "--psychology-religion",
        required=True,
        help="PMC ID for the psychology-of-religion paper (e.g. PMC11165362).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("tests/fixtures/apa_45refs/input_References.docx"),
        help="Output DOCX path.",
    )
    parser.add_argument(
        "--refs-per-paper",
        type=int,
        default=15,
        help="Number of references to extract per paper (default: 15).",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = _parse_args(argv if argv is not None else sys.argv[1:])

    api_key = os.environ.get("NCBI_API_KEY") or None
    if not api_key:
        print(
            "WARN: NCBI_API_KEY not set; falling back to anonymous "
            "E-utilities rate limit (3 req/s).",
            file=sys.stderr,
        )

    sources: list[tuple[str, str]] = [
        ("psychology", args.psychology),
        ("public-health", args.public_health),
        ("psychology-religion", args.psychology_religion),
    ]

    all_refs: list[dict[str, str]] = []
    for label, pmc_id in sources:
        try:
            xml = fetch_pmc_xml(pmc_id, api_key=api_key)
        except (urllib.error.HTTPError, urllib.error.URLError,
                TimeoutError, OSError) as e:
            print(f"ERROR: fetch_pmc_xml({pmc_id}) failed: {e}",
                  file=sys.stderr)
            return 2
        refs = extract_references(xml, limit=args.refs_per_paper)
        print(f"[{label}] extracted {len(refs)} refs from {pmc_id}")
        all_refs.extend(refs)

    build_docx(all_refs, args.output)
    print(f"wrote {len(all_refs)} refs to {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
